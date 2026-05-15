"""Populate the HM CV template with extracted data."""

import copy
import re
from datetime import date
from pathlib import Path

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from schema import HMCVData

_TEMPLATE_NAME = "HM WCM CV (sample Assistant Professor).docx"
# PyInstaller extracts everything to sys._MEIPASS (same dir as __file__ in bundle).
# In dev, the template lives one level up in assets/.
_here = Path(__file__).parent
TEMPLATE_PATH = _here / _TEMPLATE_NAME if (_here / _TEMPLATE_NAME).exists() else _here.parent / "assets" / _TEMPLATE_NAME
NA = "N/A"


def _set_para_text(para, text: str):
    """Replace all runs in a paragraph with a single run of text, preserving style."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _append_text(doc: Document, text: str, style: str = "List Paragraph"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def _find_para_containing(doc: Document, substring: str):
    for para in doc.paragraphs:
        if substring in para.text:
            return para
    return None


def _insert_after(para, new_para):
    """Insert new_para immediately after para in the document XML."""
    para._p.addnext(new_para._p)


def _add_text_after(doc: Document, anchor_para, text: str, style: str = "List Paragraph"):
    new_p = doc.add_paragraph(style=style)
    new_p.add_run(text)
    _insert_after(anchor_para, new_p)
    return new_p


def _list_items(items) -> str:
    if not items:
        return NA
    return "\n".join(f"• {item}" for item in items)


def _format_position(pos) -> str:
    dates = f" ({pos.start_year}–{pos.end_year or 'Present'})" if pos.start_year else ""
    dept = f", {pos.department}" if pos.department else ""
    return f"{pos.title}{dept}, {pos.institution}{dates}"


def _format_training(t) -> str:
    dates = f" ({t.start_year}–{t.end_year})" if t.start_year else ""
    specialty = f", {t.specialty}" if t.specialty else ""
    return f"{t.type}{specialty}, {t.institution}{dates}"


def _format_degree(d) -> str:
    field = f", {d.field}" if d.field else ""
    year = f" ({d.year})" if d.year else ""
    return f"{d.degree}{field} — {d.institution}{year}"


def _format_grant(g) -> str:
    parts = [g.source]
    if g.grant_number:
        parts.append(g.grant_number)
    parts.append(g.title)
    if g.pi:
        parts.append(f"PI: {g.pi}")
    if g.role:
        parts.append(f"Role: {g.role}")
    if g.dates:
        parts.append(g.dates)
    if g.amount:
        parts.append(g.amount)
    return " | ".join(parts)


def _format_mentee(m) -> str:
    parts = [m.name]
    if m.degree:
        parts.append(m.degree)
    if m.institution:
        parts.append(m.institution)
    if m.dates:
        parts.append(m.dates)
    if m.current_position:
        parts.append(f"→ {m.current_position}")
    return ", ".join(parts)


def _format_presentation(p) -> str:
    loc = f", {p.location}" if p.location else ""
    date_str = f". {p.date}" if p.date else ""
    return f"{p.title}. {p.venue}{loc}{date_str}"


def populate_template(data: HMCVData, output_path: str) -> str:
    doc = Document(str(TEMPLATE_PATH))
    paras = doc.paragraphs

    def find(text):
        return _find_para_containing(doc, text)

    def fill_after(anchor_text, lines: list[str], style="No Spacing"):
        anchor = find(anchor_text)
        if not anchor:
            return
        prev = anchor
        for line in (lines if lines else [NA]):
            new_p = doc.add_paragraph(style=style)
            new_p.add_run(line)
            prev._p.addnext(new_p._p)
            prev = new_p

    # ── Header ─────────────────────────────────────────────────────────────────
    name_para = find("Name:")
    if name_para:
        _set_para_text(name_para, f"Name: {data.name}")

    date_para = find("Date of Preparation:")
    if date_para:
        _set_para_text(date_para, f"Date of Preparation:  {data.date_of_preparation or date.today().strftime('%B %d, %Y')}")

    # ── Personal Data ──────────────────────────────────────────────────────────
    personal_anchor = find("PERSONAL DATA")
    if personal_anchor:
        lines = []
        if data.address:
            lines.append(f"Address: {data.address}")
        if data.phone:
            lines.append(f"Phone: {data.phone}")
        if data.email:
            lines.append(f"Email: {data.email}")
        if data.other_personal:
            lines.append(data.other_personal)
        fill_after("PERSONAL DATA", lines)

    # ── Education ──────────────────────────────────────────────────────────────
    fill_after("Academic Degree(s)", [_format_degree(d) for d in data.academic_degrees])
    fill_after("Other Educational Experiences", [e for e in data.other_educational_experiences] or [NA])

    # ── Postdoctoral Training ──────────────────────────────────────────────────
    fill_after("POSTDOCTORAL TRAINING", [_format_training(t) for t in data.postdoctoral_training])

    # ── Professional Positions ─────────────────────────────────────────────────
    fill_after("Academic Appointments", [_format_position(p) for p in data.academic_appointments])
    fill_after("Hospital Appointments", [_format_position(p) for p in data.hospital_appointments])
    fill_after("Other Professional Positions", [_format_position(p) for p in data.other_positions])

    # ── Employment Status ──────────────────────────────────────────────────────
    emp_para = find("Name of Current Employer(s):")
    if emp_para:
        _set_para_text(emp_para, f"Name of Current Employer(s):\t{data.current_employer or NA}")
    status_para = find("Current Employment Status:")
    if status_para:
        _set_para_text(status_para, f"Current Employment Status: \t{data.employment_status or NA}")

    # ── Licensure ──────────────────────────────────────────────────────────────
    lic_lines = [f"{l.state} — #{l.number or 'N/A'} (exp. {l.expiration or 'N/A'})" for l in data.licensure]
    fill_after("Licensure:", lic_lines)
    board_lines = [f"{b.board}: {b.specialty} ({b.year or 'N/A'})" for b in data.board_certifications]
    fill_after("Board Certification", board_lines)

    # ── Affiliations ───────────────────────────────────────────────────────────
    fill_after("INSTITUTIONAL/HOSPITAL AFFILIATION", data.institutional_affiliations or [NA])

    # ── Honors ─────────────────────────────────────────────────────────────────
    fill_after("HONORS, AWARDS", data.honors_awards or [NA])

    # ── Professional Organizations ─────────────────────────────────────────────
    fill_after("PROFESSIONAL ORGANIZATIONS", data.professional_organizations or [NA])

    # ── Educational Contributions ──────────────────────────────────────────────
    fill_after("Didactic teaching", data.didactic_teaching or [NA])
    fill_after("Clinical teaching", data.clinical_teaching or [NA])
    fill_after("Administrative teaching", data.administrative_teaching or [NA])
    fill_after("Continuing education", data.continuing_education or [NA])
    fill_after("Other education", data.other_education or [NA])

    # ── Clinical ───────────────────────────────────────────────────────────────
    fill_after("Clinical Practice", [data.clinical_practice or NA])
    fill_after("Clinical Innovations", [data.clinical_innovations or NA])
    fill_after("Clinical Leadership", [data.clinical_leadership or NA])

    # ── Research ───────────────────────────────────────────────────────────────
    fill_after("Research Activities:", [data.research_activities or NA])
    fill_after("Current Research Funding", [_format_grant(g) for g in data.current_funding])
    fill_after("Past (Completed) Funding", [_format_grant(g) for g in data.past_funding])
    fill_after("Pending Funding", [_format_grant(g) for g in data.pending_funding])
    fill_after("Patents & Inventions", data.patents or [NA])

    # ── Mentoring ──────────────────────────────────────────────────────────────
    fill_after("Current Mentees:", [_format_mentee(m) for m in data.current_mentees])
    fill_after("Past Mentees:", [_format_mentee(m) for m in data.past_mentees])
    fill_after("Institutional Training Grants", data.training_grants or [NA])

    # ── Leadership & Admin ─────────────────────────────────────────────────────
    fill_after("INSTITUTIONAL LEADERSHIP", data.institutional_leadership or [NA])
    fill_after("INSTITUTIONAL ADMINISTRATIVE", data.institutional_administrative or [NA])

    # ── Extramural ─────────────────────────────────────────────────────────────
    fill_after("Leadership in Extramural", data.extramural_leadership or [NA])
    fill_after("Regional", data.boards_committees_regional or [NA])
    fill_after("National", data.boards_committees_national or [NA])
    fill_after("International", data.boards_committees_international or [NA])
    fill_after("Grant Reviewing", data.grant_reviewing or [NA])
    ed_editor = [f"{e.journal} — {e.role} ({e.dates or ''})" for e in data.editorial_editor]
    fill_after("Editor/Co-Editor", ed_editor)
    ed_board = [f"{e.journal} ({e.dates or ''})" for e in data.editorial_board]
    fill_after("Editorial Board Membership", ed_board)
    ed_rev = [f"{e.journal}" for e in data.editorial_reviewer]
    fill_after("Journal Reviewing", ed_rev)

    # ── Invitations ────────────────────────────────────────────────────────────
    fill_after("Regional*", [_format_presentation(p) for p in data.presentations_regional])
    fill_after("National*", [_format_presentation(p) for p in data.presentations_national])
    fill_after("International*", [_format_presentation(p) for p in data.presentations_international])

    # ── Bibliography ───────────────────────────────────────────────────────────
    def add_pubs(anchor_text, pubs):
        lines = [p.citation for p in pubs] if pubs else [NA]
        numbered = [f"{i+1}. {line}" for i, line in enumerate(lines)] if pubs else lines
        fill_after(anchor_text, numbered, style="List Paragraph")

    add_pubs("Peer-reviewed Research Articles:", data.peer_reviewed_articles)
    add_pubs("Reviews and Editorials:", data.reviews_editorials)
    add_pubs("Books:", data.books)
    add_pubs("Chapters:", data.chapters)
    add_pubs("Non-peer-reviewed Research Publications:", data.non_peer_reviewed)
    add_pubs("Case Reports", data.case_reports)
    add_pubs("In review", data.in_review)
    add_pubs("Abstracts", data.abstracts)
    add_pubs("Other (media, podcasts", data.other_publications)

    doc.save(output_path)
    return output_path
