"""Extract structured CV data from any resume using DeepSeek API."""

import json
from pathlib import Path

import docx
from openai import OpenAI
from pdfminer.high_level import extract_text as pdf_extract_text

import config
from schema import HMCVData


def _client() -> OpenAI:
    key = config.get_api_key()
    if not key:
        raise RuntimeError("DeepSeek API key not set. Open Settings to add your key.")
    return OpenAI(api_key=key, base_url="https://api.deepseek.com")

SYSTEM_PROMPT = """You are a medical CV data extraction specialist. Extract ALL information from the provided resume/CV text and return a single valid JSON object matching the schema exactly.

Rules:
- Extract every piece of information present. Do not invent or fabricate data.
- If a field has no data in the source document, return an empty list [] or null.
- For publications: include the COMPLETE citation string. Bold the candidate's name with **name** markdown.
- For dates: use the format found in the source (e.g. "2018-2022", "June 2018 - Present").
- For grants: extract source agency, grant number if visible, title, PI, role, dates, dollar amount if listed.
- For the name field: include all credentials (MD, PhD, etc.) if present.
- Return ONLY the JSON object. No explanation, no markdown code fences, just raw JSON."""

EXTRACTION_SCHEMA_HINT = json.dumps(HMCVData.model_json_schema(), indent=2)


def read_docx(path: str) -> str:
    doc = docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_pdf(path: str) -> str:
    return pdf_extract_text(path)


def read_resume(path: str) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in (".docx",):
        return read_docx(path)
    elif suffix == ".pdf":
        return read_pdf(path)
    elif suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def extract_cv_data(resume_path: str) -> HMCVData:
    """Extract structured HM CV data from any resume file."""
    resume_text = read_resume(resume_path)

    user_message = f"""Extract all CV/resume data from the following document. Return JSON matching this schema:

{EXTRACTION_SCHEMA_HINT}

Document to extract from:
---
{resume_text}
---"""

    response = _client().chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_message},
        ],
        temperature=0.1,
        max_tokens=8000,
        response_format={"type": "json_object"},
    )

    raw_json = response.choices[0].message.content
    data = json.loads(raw_json)
    return HMCVData.model_validate(data)
